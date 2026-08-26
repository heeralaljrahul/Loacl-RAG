"""Turn a file on disk into plain text blocks.

A loader returns a list of ``Block``s.  A block is a run of text plus the
location it came from (a page number for PDFs, nothing for flat text).
Keeping page numbers here is what lets an answer cite "handbook.pdf p.14"
instead of just "handbook.pdf".
"""

from __future__ import annotations

import csv
import io
import json
import re
from dataclasses import dataclass
from pathlib import Path

TEXT_EXTS = {
    ".txt", ".md", ".markdown", ".rst", ".org", ".tex", ".log",
    ".py", ".js", ".jsx", ".ts", ".tsx", ".java", ".kt", ".go", ".rs",
    ".c", ".h", ".cpp", ".hpp", ".cs", ".rb", ".php", ".swift", ".lua",
    ".sh", ".bat", ".ps1", ".sql", ".yaml", ".yml", ".toml", ".ini", ".cfg",
}
DOC_EXTS = {".pdf", ".docx", ".html", ".htm", ".json", ".jsonl", ".csv", ".tsv", ".epub"}
SUPPORTED_EXTS = TEXT_EXTS | DOC_EXTS


@dataclass
class Block:
    text: str
    page: int | None = None


class LoaderError(RuntimeError):
    pass


def is_supported(path: Path) -> bool:
    return path.suffix.lower() in SUPPORTED_EXTS


def load(path: Path) -> list[Block]:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return _load_pdf(path)
    if ext == ".docx":
        return _load_docx(path)
    if ext == ".epub":
        return _load_epub(path)
    if ext in (".html", ".htm"):
        return _load_html(path)
    if ext in (".json", ".jsonl"):
        return _load_json(path)
    if ext in (".csv", ".tsv"):
        return _load_csv(path)
    return _load_text(path)


# --------------------------------------------------------------------------


def _read_text(path: Path) -> str:
    raw = path.read_bytes()
    for enc in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def _load_text(path: Path) -> list[Block]:
    return [Block(_read_text(path))]


def _load_pdf(path: Path) -> list[Block]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover
        raise LoaderError("pypdf is required for PDFs: pip install pypdf") from exc

    reader = PdfReader(str(path))
    if getattr(reader, "is_encrypted", False):
        try:
            reader.decrypt("")
        except Exception as exc:
            raise LoaderError("PDF is password protected") from exc

    blocks: list[Block] = []
    for i, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        text = _fix_pdf_text(text)
        if text.strip():
            blocks.append(Block(text, page=i))
    if not blocks:
        raise LoaderError(
            "no extractable text (scanned PDF?) — OCR it first, e.g. with ocrmypdf"
        )
    return blocks


def _fix_pdf_text(text: str) -> str:
    # PDF extraction hard-wraps mid-sentence and hyphenates across lines.
    text = re.sub(r"-\n(\w)", r"\1", text)
    text = re.sub(r"(?<![.!?:;\n])\n(?![\n\s•\-*\d])", " ", text)
    return re.sub(r"[ \t]{2,}", " ", text)


def _load_docx(path: Path) -> list[Block]:
    try:
        import docx  # python-docx
    except ImportError as exc:  # pragma: no cover
        raise LoaderError("python-docx is required for .docx files") from exc

    document = docx.Document(str(path))
    parts: list[str] = []
    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            parts.append("")
            continue
        style = (para.style.name or "").lower() if para.style else ""
        if style.startswith("heading"):
            level = "".join(ch for ch in style if ch.isdigit()) or "1"
            parts.append("")
            parts.append("#" * min(int(level), 6) + " " + text)
            parts.append("")
        else:
            parts.append(text)
    for table in document.tables:
        rows = [
            " | ".join(cell.text.strip().replace("\n", " ") for cell in row.cells)
            for row in table.rows
        ]
        rows = [r for r in rows if r.strip(" |")]
        if rows:
            parts += ["", *rows, ""]
    return [Block("\n".join(parts))]


def _load_epub(path: Path) -> list[Block]:
    import zipfile

    blocks: list[Block] = []
    with zipfile.ZipFile(path) as zf:
        names = sorted(
            n for n in zf.namelist()
            if n.lower().endswith((".xhtml", ".html", ".htm"))
        )
        for name in names:
            text = _html_to_text(zf.read(name).decode("utf-8", errors="replace"))
            if text.strip():
                blocks.append(Block(text))
    if not blocks:
        raise LoaderError("no readable chapters found in epub")
    return blocks


def _load_html(path: Path) -> list[Block]:
    return [Block(_html_to_text(_read_text(path)))]


def _html_to_text(html: str) -> str:
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        text = re.sub(r"(?is)<(script|style)[^>]*>.*?</\1>", " ", html)
        text = re.sub(r"(?i)<br\s*/?>|</p>|</div>|</li>", "\n", text)
        text = re.sub(r"<[^>]+>", " ", text)
        return _collapse(text)

    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "noscript"]):
        tag.decompose()
    out: list[str] = []
    for el in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "pre", "td"]):
        text = el.get_text(" ", strip=True)
        if not text:
            continue
        if el.name.startswith("h") and len(el.name) == 2:
            out += ["", "#" * int(el.name[1]) + " " + text, ""]
        elif el.name == "li":
            out.append("- " + text)
        else:
            out.append(text)
    return _collapse("\n".join(out)) if out else _collapse(soup.get_text(" "))


def _load_json(path: Path) -> list[Block]:
    text = _read_text(path)
    if path.suffix.lower() == ".jsonl":
        lines = []
        for line in text.splitlines():
            line = line.strip()
            if not line:
                continue
            try:
                lines.append(_flatten_json(json.loads(line)))
            except json.JSONDecodeError:
                lines.append(line)
        return [Block("\n\n".join(lines))]
    try:
        return [Block(_flatten_json(json.loads(text)))]
    except json.JSONDecodeError:
        return [Block(text)]


def _flatten_json(obj, prefix: str = "") -> str:
    """Render JSON as ``path: value`` lines — keeps keys next to their values
    so both the keyword and the vector side have something to match on."""
    lines: list[str] = []
    if isinstance(obj, dict):
        for key, value in obj.items():
            lines.append(_flatten_json(value, f"{prefix}.{key}" if prefix else str(key)))
    elif isinstance(obj, list):
        for i, value in enumerate(obj):
            lines.append(_flatten_json(value, f"{prefix}[{i}]"))
    else:
        lines.append(f"{prefix}: {obj}" if prefix else str(obj))
    return "\n".join(l for l in lines if l)


def _load_csv(path: Path) -> list[Block]:
    delim = "\t" if path.suffix.lower() == ".tsv" else ","
    text = _read_text(path)
    reader = csv.reader(io.StringIO(text), delimiter=delim)
    rows = list(reader)
    if not rows:
        return [Block("")]
    header, *body = rows
    # One block per 40 rows, header repeated, so a chunk is never orphaned
    # from its column names.
    blocks: list[Block] = []
    for start in range(0, len(body), 40):
        window = body[start:start + 40]
        lines = [" | ".join(header)]
        lines += [" | ".join(cell.replace("\n", " ") for cell in row) for row in window]
        blocks.append(Block("\n".join(lines)))
    return blocks or [Block(" | ".join(header))]


def _collapse(text: str) -> str:
    text = text.replace("\xa0", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return "\n".join(line.strip() for line in text.splitlines()).strip()
